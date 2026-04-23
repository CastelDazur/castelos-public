"""
text_quality_checker.py - Post-generation quality audit for Text Studio outputs.

Checks a text file or docx against CastelOS anti-slop rules and content requirements.
Returns a structured report with pass/fail per check and a publish-ready verdict.

Usage:
    python text_quality_checker.py input.txt
    python text_quality_checker.py input.docx --keyword "local llm workstation 2026"
    python text_quality_checker.py input.docx --keyword "cold outreach saas" --max-words 1200
"""

import re
import sys
import argparse
from dataclasses import dataclass, field
from pathlib import Path


BLACKLIST_WORDS = [
    "leverage", "leveraging", "robust", "cutting-edge", "cutting edge",
    "transformative", "unlock", "seamless", "seamlessly",
    "game-changing", "game-changer", "revolutionary", "revolutionize",
    "navigate", "navigating", "empower", "foster", "streamline",
    "synergy", "holistic", "delve into", "dive into", "harness",
    "enable you to", "bespoke", "curated", "ecosystem", "paradigm",
    "tapestry", "journey", "realm", "landscape", "embark",
]

BANNED_PHRASES = [
    "matters more than people realize",
    "it is worth noting",
    "it's worth noting",
    "when it comes to",
    "at the end of the day",
    "in today's world",
    "in the ever-evolving",
]


@dataclass
class CheckResult:
    name: str
    passed: bool
    detail: str
    severity: str = "high"


@dataclass
class AuditReport:
    checks: list = field(default_factory=list)
    word_count: int = 0
    source_file: str = ""

    @property
    def passed(self) -> bool:
        return all(c.passed for c in self.checks if c.severity == "high")

    @property
    def high_failures(self) -> list:
        return [c for c in self.checks if not c.passed and c.severity == "high"]

    def summary(self) -> str:
        lines = []
        lines.append(f"File: {self.source_file}")
        lines.append(f"Words: {self.word_count}")
        lines.append(f"Checks: {len(self.checks)}")
        lines.append("")
        for c in self.checks:
            icon = "PASS" if c.passed else "FAIL"
            lines.append(f"  [{icon}] {c.name} ({c.severity}): {c.detail}")
        lines.append("")
        if self.passed:
            lines.append("VERDICT: Publishable (all high-severity checks passed)")
        else:
            fails = self.high_failures
            lines.append(f"VERDICT: Not publishable ({len(fails)} high-severity failure(s))")
            for f in fails:
                lines.append(f"  - {f.name}: {f.detail}")
        return "\n".join(lines)


def load_text(path):
    if path.suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            print("python-docx required: pip install python-docx")
            sys.exit(1)
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    return path.read_text(encoding="utf-8")


def check_em_dash(text):
    count = text.count("\u2014")
    return CheckResult("Em-dash", count == 0, f"{count} found" if count else "0 found")


def check_blacklist(text):
    hits = []
    for w in BLACKLIST_WORDS:
        c = len(re.findall(r"\\b" + re.escape(w) + r"\\b", text, re.IGNORECASE))
        if c: hits.append((w, c))
    if not hits:
        return CheckResult("Blacklist words", True, "0 found")
    return CheckResult("Blacklist words", False, str(sum(c for _,c in hits)) + " found")


def check_citation_integrity(text):
    refs = set(int(n) for n in re.findall(r"\\[source:(\\d+)\\]", text))
    if not refs:
        return CheckResult("Citation integrity", True, "no refs", severity="low")
    defined = set()
    for m in re.finditer(r"^\\s*\\[(\\d+)\\]\\s+\\S", text, re.MULTILINE):
        defined.add(int(m.group(1)))
    missing = refs - defined
    if not missing:
        return CheckResult("Citation integrity", True, str(len(refs)) + " refs matched")
    return CheckResult("Citation integrity", False, "refs " + str(sorted(missing)) + " missing")


def check_duplicates(text):
    sentences = [s.strip() for s in re.split(r"[.!?]\\s+", text) if len(s.strip()) > 50]
    from collections import Counter
    dups = [(s,c) for s,c in Counter(sentences).most_common(5) if c > 1]
    if not dups:
        return CheckResult("Duplicate sentences", True, "none found")
    return CheckResult("Duplicate sentences", False, str(len(dups)) + " duplicate(s)")


def audit(text, keyword="", min_words=800, max_words=3000, source_file=""):
    report = AuditReport(source_file=source_file, word_count=len(text.split()))
    report.checks.append(check_em_dash(text))
    report.checks.append(check_blacklist(text))
    report.checks.append(check_citation_integrity(text))
    report.checks.append(check_duplicates(text))
    return report


def main():
    parser = argparse.ArgumentParser(description="CastelOS Text Quality Checker")
    parser.add_argument("file", help="Input file (.txt or .docx)")
    parser.add_argument("--keyword", default="", help="Target SEO keyword")
    parser.add_argument("--max-words", type=int, default=3000)
    args = parser.parse_args()
    path = Path(args.file)
    if not path.exists():
        print(f"File not found: {path}")
        sys.exit(1)
    text = load_text(path)
    report = audit(text, keyword=args.keyword, max_words=args.max_words, source_file=str(path))
    print(report.summary())
    sys.exit(0 if report.passed else 1)


if __name__ == "__main__":
    main()
